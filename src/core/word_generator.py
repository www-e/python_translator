from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.logger import setup_logger

logger = setup_logger(__name__)

class WordDocumentGenerator:
    """Generate Word (.docx) document with proper Arabic/English BiDi support"""
    
    def __init__(self, output_path):
        self.output_path = output_path
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Setup document settings"""
        # Set document margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def _set_paragraph_bidi(self, paragraph):
        """Set BiDi (bidirectional) property for paragraph to handle mixed Arabic/English text"""
        # Access the paragraph's XML element
        pPr = paragraph._element.get_or_add_pPr()
        
        # Add BiDi element (bidi = bidirectional text)
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')  # Enable BiDi
        pPr.append(bidi)
        
        # Set right-to-left alignment
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _set_run_rtl(self, run):
        """Set RTL (right-to-left) properties for text run"""
        # Set complex script font (required for Arabic)
        run.font.complex_script = True
        run.font.rtl = True
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    
    def generate_document(self, translated_pages, progress_callback=None):
        """Generate Word document with proper BiDi Arabic text"""
        logger.info(f"Generating Word document: {self.output_path}")
        
        try:
            total_pages = len(translated_pages)
            
            for i, page_data in enumerate(translated_pages):
                self._add_page_content(page_data['text'])
                
                if progress_callback:
                    progress_callback(i + 1, total_pages)
            
            # Save document
            self.doc.save(self.output_path)
            logger.info(f"Word document generated successfully: {self.output_path}")
            
        except Exception as e:
            logger.error(f"Word document generation failed: {str(e)}")
            raise
    
    def _add_page_content(self, text):
        """Add page content with proper BiDi support"""
        # Split into paragraphs
        paragraphs_text = text.split('\n')
        
        for para_text in paragraphs_text:
            if para_text.strip():
                # Add paragraph
                para = self.doc.add_paragraph()
                
                # Set BiDi property for mixed Arabic/English
                self._set_paragraph_bidi(para)
                
                # Add text as run
                run = para.add_run(para_text.strip())
                
                # Set RTL properties
                self._set_run_rtl(run)
            else:
                # Add empty paragraph for spacing
                self.doc.add_paragraph()
